"""Document processing service for RAG functionality."""

import logging
import os
import re
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentService:
    """Handles document upload processing, text extraction, and RAG context building."""

    @staticmethod
    def process_file(filepath: str, file_type: str) -> Optional[str]:
        """Extract text from a file based on its type.

        Returns:
            Extracted text or None on failure.
        """
        processors = {
            "pdf": DocumentService._process_pdf,
            "docx": DocumentService._process_docx,
            "txt": DocumentService._process_txt,
        }

        processor = processors.get(file_type.lower())
        if not processor:
            logger.error("Unsupported file type: %s", file_type)
            return None

        try:
            text = processor(filepath)
            if text:
                # Clean up extracted text
                text = re.sub(r"\n{3,}", "\n\n", text)  # Remove excessive newlines
                text = text.strip()
            return text
        except Exception as e:
            logger.exception("Error processing file %s: %s", filepath, e)
            return None

    @staticmethod
    def _process_pdf(filepath: str) -> Optional[str]:
        """Extract text from a PDF file."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(filepath)
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {i + 1}]\n{page_text}")

            return "\n\n".join(text_parts) if text_parts else None
        except Exception as e:
            logger.error("PDF processing error: %s", e)
            return None

    @staticmethod
    def _process_docx(filepath: str) -> Optional[str]:
        """Extract text from a DOCX file."""
        try:
            from docx import Document

            doc = Document(filepath)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts) if text_parts else None
        except Exception as e:
            logger.error("DOCX processing error: %s", e)
            return None

    @staticmethod
    def _process_txt(filepath: str) -> Optional[str]:
        """Read text from a plain text file."""
        try:
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
            for encoding in encodings:
                try:
                    with open(filepath, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return None
        except Exception as e:
            logger.error("TXT processing error: %s", e)
            return None

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
        """Split text into overlapping chunks for retrieval.

        Args:
            text: The full document text.
            chunk_size: Maximum characters per chunk.
            overlap: Number of overlapping characters between chunks.
        """
        if not text:
            return []

        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size

            # Try to break at a sentence boundary
            if end < len(text):
                # Look for sentence end near the chunk boundary
                for sep in [". ", ".\n", "\n\n", "\n", " "]:
                    boundary = text.rfind(sep, start + chunk_size // 2, end)
                    if boundary != -1:
                        end = boundary + len(sep)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap
            if start >= len(text):
                break

        return chunks

    @staticmethod
    def search_chunks(query: str, chunks: list[str], top_k: int = 3) -> list[str]:
        """Simple keyword-based relevance search over text chunks.

        Returns the top_k most relevant chunks based on keyword overlap.
        """
        if not query or not chunks:
            return []

        # Tokenize query
        query_words = set(re.findall(r"\w+", query.lower()))
        if not query_words:
            return chunks[:top_k]

        scored = []
        for chunk in chunks:
            chunk_lower = chunk.lower()
            # Score based on keyword matches
            score = sum(1 for word in query_words if word in chunk_lower)
            # Bonus for exact phrase match
            if query.lower() in chunk_lower:
                score += 5
            scored.append((score, chunk))

        # Sort by score descending
        scored.sort(key=lambda x: x[0], reverse=True)

        # Return top_k chunks that have at least some relevance
        results = [chunk for score, chunk in scored[:top_k] if score > 0]
        return results if results else chunks[:top_k]  # Fallback to first chunks

    @classmethod
    def build_rag_context(cls, query: str, documents_text: list[tuple[str, str]], max_context_chars: int = 4000) -> str:
        """Build context from multiple documents for RAG.

        Args:
            query: The user's question.
            documents_text: List of (filename, extracted_text) tuples.
            max_context_chars: Maximum characters for the context.

        Returns:
            Formatted context string for injection into the AI prompt.
        """
        if not documents_text:
            return ""

        all_results = []
        for filename, text in documents_text:
            if not text:
                continue
            chunks = cls.chunk_text(text)
            relevant = cls.search_chunks(query, chunks, top_k=2)
            for chunk in relevant:
                all_results.append(f"[Source: {filename}]\n{chunk}")

        if not all_results:
            return ""

        # Combine results, respecting max context size
        context_parts = []
        total_chars = 0
        for result in all_results:
            if total_chars + len(result) > max_context_chars:
                break
            context_parts.append(result)
            total_chars += len(result)

        return "\n\n---\n\n".join(context_parts)

    @staticmethod
    def allowed_file(filename: str, allowed_extensions: set) -> bool:
        """Check if a filename has an allowed extension."""
        return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_extensions

    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get the file extension without the dot."""
        return filename.rsplit(".", 1)[1].lower() if "." in filename else ""

    @staticmethod
    def get_file_size(filepath: str) -> int:
        """Get file size in bytes."""
        try:
            return os.path.getsize(filepath)
        except OSError:
            return 0
